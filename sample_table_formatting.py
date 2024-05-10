from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Initialize the Document
doc = Document()

# Create a table
table = doc.add_table(rows=4, cols=5)  # Adjust the number of rows as needed

# Define the header cells
header_cells = table.rows[0].cells
header_texts = [
    "Population",
    "Parameter",
    "Current Threshold",
    "BTL Threshold",
    "Recommended Threshold",
]

for cell, text in zip(header_cells, header_texts):
    cell.text = text
    cell.vertical_alignment = (
        WD_ALIGN_VERTICAL.BOTTOM
    )  # Align text to bottom (optional)
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.font.bold = True  # Make header text bold

# Add rows to the table (example data)
data_rows = [
    ["Business Non-High", "Minimum Value", "$5,000.00", "$3,700.00", "$5,000.00"],
    ["Business Non-High", "Minimum Volume", "3", "2", "3"],
    ["Business Non-High", "No. of Occurrences", "1", "1", "1"],
]

for row_data, row in zip(data_rows, table.rows[1:]):
    for cell, text in zip(row.cells, row_data):
        cell.text = text

# Function to set cell borders
def set_cell_borders(cell, border_type, sz=6, color="auto"):
    """
    Set cell borders. Specify the border_type as 'top', 'bottom', 'start', or 'end'.
    Set 'sz' to the size of the border and 'color' to the border color.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tag = "w:{}".format(border_type)
    element = tcPr.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        tcPr.append(element)
    element.set(qn("w:sz"), str(sz))
    element.set(qn("w:val"), "single")
    element.set(qn("w:color"), color)


# Apply formatting to the table
for row in table.rows:
    for cell in row.cells:
        # Clear all cell borders
        set_cell_borders(cell, "top", sz=0, color="FFFFFF")
        set_cell_borders(cell, "bottom", sz=0, color="FFFFFF")
        set_cell_borders(cell, "start", sz=0, color="FFFFFF")
        set_cell_borders(cell, "end", sz=0, color="FFFFFF")

# Set top and bottom borders for header row
for cell in table.rows[0].cells:
    set_cell_borders(cell, "top", sz=6, color="000000")
    set_cell_borders(cell, "bottom", sz=6, color="000000")

# Set bottom border for last row
for cell in table.rows[-1].cells:
    set_cell_borders(cell, "bottom", sz=6, color="000000")

# Save the document
doc.save("formatted_table.docx")
