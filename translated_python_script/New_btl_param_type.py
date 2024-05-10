# TODO:
import docx
import numpy as np
import os
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Pt
from docx.shared import Pt
from docx.shared import RGBColor
from docx.text.paragraph import WD_STYLE_TYPE
from natsort import natsorted

# Specify location of the tuning tracker
fp = r"C:/Users/KadamatiV/OneDrive - Crowe LLP/Documents/PROJECTHUB/TM TUNING/TM_Narrative_Codebase_Consolidation/required_processing_data/"
file_name = "BTL Tuning Tracker - With Calculations.xlsx"

# Specify how to format the different parameters seen (change threshold names as needed)
currencyF = [
    "Minimal Sum",
    "Minimum Value",
    "Minimal Transaction Amount",
    "Sum Lower Bound",
    "Minimal Current Month Sum",
    "Minimal Transaction Value",
    "Transaction Amount Lower Bound",
    "Sum Amount Lower Bound",
]
numberF = ["No. of Occurrences", "Minimum Volume", "Min Value"]
percentF = ["Ratio Lower Bound", "Ratio Upper Bound"]
decimalF = [
    "STDEV exceeds Historical Average Sum",
    "STDEV exceeds Historical Average Count",
]

# Data is loaded
data = pd.read_excel(os.path.join(fp, file_name), sheet_name=0)

# Enter the column name of the Rule IDs, if different
ruleIDs = data["Rule ID"].value_counts().index.tolist()

# Create lookup for values below 10 that need to be written out (no change)
numbers = ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]
alpha_numbers = [
    "zero (0)",
    "one (1)",
    "two (2)",
    "three (3)",
    "four (4)",
    "five (5)",
    "six (6)",
    "seven (7)",
    "eight (8)",
    "nine (9)",
]
alpha_numbers_cap = [
    "Zero (0)",
    "One (1)",
    "Two (2)",
    "Three (3)",
    "Four (4)",
    "Five (5)",
    "Six (6)",
    "Seven (7)",
    "Eight (8)",
    "Nine (9)",
]
numbers_df = pd.DataFrame(
    {
        "numbers": numbers,
        "alpha_numbers": alpha_numbers,
        "alpha_numbers_cap": alpha_numbers_cap,
    }
)

# Create an empty data frame to hold narratives
# narratives = pd.DataFrame(
#     columns=[
#         "Rule ID",
#         "Population Group",
#         "Summary",
#         "Analysis",
#         "Conclusion",
#     ]
# )
# TODO Change log for above create an empty data frame to hold narratives
# Create an empty DataFrame with 5 columns
narratives = pd.DataFrame(columns=range(5))

# Begin Iterations --------------------------------------------------------
# Iterate over Rule IDs, Population Groups, and Parameters
for x in ruleIDs:
    data_Rule = data[data["Rule ID"] == x]
    popGroups = data_Rule["Population Group"].value_counts().index.tolist()

    for pop in popGroups:
        data_Rule_Pop = data_Rule[data_Rule["Population Group"] == pop]
        param_types = data_Rule_Pop["Parameter Type"].value_counts().index.tolist()

        # Parse date range of alert generation for the rule
        date_range = data_Rule_Pop["Date Range"].iloc[0].split("-")

        for param_type in param_types:
            data_Rule_Pop_Param_Type = data_Rule_Pop[
                data_Rule_Pop["Parameter Type"] == param_type
                ]

            parameter = data_Rule_Pop_Param_Type["Parameter"].iloc[0]
            param_type_value = data_Rule_Pop_Param_Type["Parameter Type"].iloc[0]

            # Summary Paragraph -----------------------------------------------------------------
            # Time frame of alerts that generated; if no alerts generated -> no analysis performed
            temp = data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
            date_range = data_Rule_Pop_Param_Type["Date Range"].iloc[0].split("-")

            line_one = ""
            if temp == 0:
                line_one = f"No alerts generated in the Actimize UAT environment between {date_range[0]} and {date_range[1]}; therefore, no analysis was performed. The current thresholds are recommended to be maintained."
            elif temp > 0:
                line_one = f"Alerts generated in the Actimize UAT environment between {date_range[0]} and {date_range[1]} were extracted for review."

            # No. of rule breaks generated; wording reflects if there was sampling or data quality alerts present
            temp = data_Rule_Pop_Param_Type["Num Alerts Sampled"].iloc[0]

            line_two = ""
            if temp == 0:
                pass
            elif (
                    temp == 1
                    and temp == data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
            ):
                line_two = f"{numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule break generated during the testing period, which was reviewed by the Bank for quality."
            elif (
                    temp < 10
                    and temp == data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
                    and data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0] == 0
            ):
                line_two = f"{numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule breaks generated during the testing period, which were reviewed by the Bank for quality and used for analysis."
            elif (
                    temp < 10
                    and temp == data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
            ):
                line_two = f"{numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule breaks generated during the testing period, which were reviewed by the Bank for quality."
            elif (
                    temp == data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
                    and data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0] == 0
            ):
                line_two = f"{data_Rule_Pop_Param_Type['Num Alerts Extracted'].iloc[0]} rule breaks generated during the testing period, which were reviewed by the Bank for quality and used for analysis."
            elif (
                    temp == data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
                    and data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0] != 0
            ):
                line_two = f"{data_Rule_Pop_Param_Type['Num Alerts Extracted'].iloc[0]} rule breaks generated during the testing period, which were reviewed by the Bank for quality."
            elif (
                    temp != data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]
                    and data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0] == 0
            ):
                line_two = f"{data_Rule_Pop_Param_Type['Num Alerts Extracted'].iloc[0]:,} rule breaks generated during the testing period, which were subsequently sampled using the standard methodology. The {temp} sampled rule breaks were then reviewed by the Bank for quality and used for analysis."
            elif temp != data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0]:
                line_two = f"{data_Rule_Pop_Param_Type['Num Alerts Extracted'].iloc[0]:,} rule breaks generated during the testing period, which were subsequently sampled using the standard methodology. The {temp} sampled rule breaks were then reviewed by the Bank for quality."

            # Data quality alerts identified, if any
            temp = data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0]

            line_three = ""
            if temp == 0:
                pass
            elif temp == 1:
                line_three = f"The Bank identified {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers'].iloc[0]} Data Quality rule break that generated on ###INSERT DATA QUALITY RATIONALE###. As a result, this rule break was marked as Data Quality and excluded from analysis."
            elif temp < 10:
                line_three = f"The Bank identified {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers'].iloc[0]} Data Quality rule breaks that generated on ###INSERT DATA QUALITY RATIONALE###. As a result, these rule breaks were marked as Data Quality and excluded from analysis."
            elif temp >= 10:
                line_three = f"The Bank identified {temp} Data Quality rule breaks that generated on ###INSERT DATA QUALITY RATIONALE###. As a result, these rule breaks were marked as Data Quality and excluded from analysis."

            # Reflects sampled rules used for analysis (excluding data quality)
            temp = (
                    data_Rule_Pop_Param_Type["Num Alerts Sampled"].iloc[0]
                    - data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0]
            )

            line_four = ""
            if temp == 0 or data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0] == 0:
                pass
            elif (
                    temp < 10
                    and temp != data_Rule_Pop_Param_Type["Num Alerts Sampled"].iloc[0]
            ):
                line_four = f"The remaining {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers'].iloc[0]} sampled rule breaks were used for analysis."
            elif (
                    temp >= 10
                    and temp != data_Rule_Pop_Param_Type["Num Alerts Sampled"].iloc[0]
            ):
                line_four = (
                    f"The remaining {temp} sampled rule breaks were used for analysis."
                )
            elif (
                    temp < 10
                    and data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] <= 100
            ):
                line_four = f"{numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule breaks were used for analysis."
            elif (
                    temp >= 10
                    and data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] <= 100
            ):
                line_four = f"{temp} rule breaks were used for analysis."
            elif temp < 10:
                line_four = f"{numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} sampled rule breaks were used for analysis."
            elif temp >= 10:
                line_four = f"{temp} sampled rule breaks were used for analysis."

            # Analysis Paragraph ----------------------------------------------------------------
            # Interrupt loop if parameter is not tunable
            is_tunable = data_Rule_Pop_Param_Type["Is Tunable"].iloc[0]
            if is_tunable == "No":
                continue

            # Threshold tuned and value
            temp_lowered = data_Rule_Pop_Param_Type["BTL Threshold"].iloc[0]

            if param_type_value in currencyF:
                temp_lowered_formatted = f"${temp_lowered:,.2f}"
            elif param_type_value in numberF:
                if temp_lowered < 10:
                    temp_lowered_formatted = numbers_df[
                        numbers_df["numbers"] == str(int(temp_lowered))
                        ]["alpha_numbers"].iloc[0]
                else:
                    temp_lowered_formatted = f"{temp_lowered:,}"
            elif param_type_value in percentF:
                temp_lowered_formatted = f"{temp_lowered:.7f}%"
            elif param_type_value in decimalF:
                temp_lowered_formatted = f"{temp_lowered:.7f}"
            else:
                temp_lowered_formatted = str(temp_lowered)

            temp_original = data_Rule_Pop_Param_Type["Current Threshold"].iloc[0]

            if param_type_value in currencyF:
                temp_original_formatted = f"${temp_original:,.2f}"
            elif param_type_value in numberF:
                if temp_original < 10:
                    temp_original_formatted = numbers_df[
                        numbers_df["numbers"] == str(int(temp_original))
                        ]["alpha_numbers"].iloc[0]
                else:
                    temp_original_formatted = f"{temp_original:,}"
            elif param_type_value in percentF:
                temp_original_formatted = f"{temp_original:.7f}%"
            elif param_type_value in decimalF:
                temp_original_formatted = f"{temp_original:.7f}"
            else:
                temp_original_formatted = str(temp_original)

            line_five = ""
            if data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] == 0:
                pass
            elif data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] > 0:
                line_five = f"Production below-the-line calibration was conducted on the {parameter} threshold, which was lowered from the production value of {temp_original_formatted} to the below-the-line testing value of {temp_lowered_formatted}."

            # Values that rule breaks generated at
            temp = (
                    data_Rule_Pop_Param_Type["Max Val"].iloc[0]
                    - data_Rule_Pop_Param_Type["Min Val"].iloc[0]
            )

            if param_type_value in currencyF:
                val_formatted = (
                    f"${data_Rule_Pop_Param_Type['Min Val'].iloc[0]:,.2f}",
                    f"${data_Rule_Pop_Param_Type['Max Val'].iloc[0]:,.2f}",
                )
            elif param_type_value in numberF:
                if np.isnan(data_Rule_Pop_Param_Type["Max Val"].iloc[0]):
                    val_formatted = ("0", "0")
                elif data_Rule_Pop_Param_Type["Max Val"].iloc[0] < 10:
                    val_formatted = (
                        numbers_df[
                            numbers_df["numbers"]
                            == str(int(data_Rule_Pop_Param_Type["Min Val"].iloc[0]))
                            ]["alpha_numbers"].iloc[0],
                        numbers_df[
                            numbers_df["numbers"]
                            == str(int(data_Rule_Pop_Param_Type["Max Val"].iloc[0]))
                            ]["alpha_numbers"].iloc[0],
                    )
                elif (
                        data_Rule_Pop_Param_Type["Max Val"].iloc[0] >= 10
                        and data_Rule_Pop_Param_Type["Min Val"].iloc[0] < 10
                ):
                    val_formatted = (
                        numbers_df[
                            numbers_df["numbers"]
                            == str(int(data_Rule_Pop_Param_Type["Min Val"].iloc[0]))
                            ]["alpha_numbers"].iloc[0],
                        f"{data_Rule_Pop_Param_Type['Max Val'].iloc[0]:,}",
                    )
                else:
                    val_formatted = (
                        f"{data_Rule_Pop_Param_Type['Min Val'].iloc[0]:,}",
                        f"{data_Rule_Pop_Param_Type['Max Val'].iloc[0]:,}",
                    )
            elif param_type_value in percentF:
                val_formatted = (
                    f"{data_Rule_Pop_Param_Type['Min Val'].iloc[0]:.7f}%",
                    f"{data_Rule_Pop_Param_Type['Max Val'].iloc[0]:.7f}%",
                )
            elif param_type_value in decimalF:
                val_formatted = (
                    f"{data_Rule_Pop_Param_Type['Min Val'].iloc[0]:.7f}",
                    f"{data_Rule_Pop_Param_Type['Max Val'].iloc[0]:.7f}",
                )
            else:
                val_formatted = (
                    str(data_Rule_Pop_Param_Type["Min Val"].iloc[0]),
                    str(data_Rule_Pop_Param_Type["Max Val"].iloc[0]),
                )


            # -------------- Phase 2 -------------
            line_six = ""
            if data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] == 0:
                pass
            elif temp == 0:
                line_six = f"Rule breaks were generated solely at a value of {val_formatted[1]}."
            elif temp != 0:
                line_six = f"Rule breaks were generated for values ranging between {val_formatted[0]} and {val_formatted[1]}."

            # No. of interesting rule breaks in sample population
            temp = (
                    data_Rule_Pop_Param_Type["Num Alerts Sampled"].iloc[0]
                    - data_Rule_Pop_Param_Type["Data Quality Alerts"].iloc[0]
            )
            temp2 = data_Rule_Pop_Param_Type["Interesting Alerts"].iloc[0]

            line_seven = ""
            if data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] == 0:
                pass
            elif temp2 == 0 and temp == 1:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule breaks in the population of {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule break."
            elif temp2 == 0 and temp < 10:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule breaks in the population of {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule breaks."
            elif temp2 == 0 and temp >= 10:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule breaks in the population of {temp} rule breaks."
            elif temp2 == 1 and temp == 1:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule break in the population of {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule break."
            elif temp2 == 1 and temp < 10:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule break in the population of {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule breaks."
            elif temp2 == 1 and temp >= 10:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule break in the population of {temp} rule breaks."
            elif temp2 < 10 and temp < 10:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule breaks in the population of {numbers_df[numbers_df['numbers'] == str(temp)]['alpha_numbers_cap'].iloc[0]} rule breaks."
            elif temp2 < 10 and temp >= 10:
                line_seven = f"Analysis revealed {numbers_df[numbers_df['numbers'] == str(temp2)]['alpha_numbers'].iloc[0]} interesting rule breaks in the population of {temp} rule breaks."
            elif temp2 >= 10:
                line_seven = f"Analysis revealed {temp2} interesting rule breaks in the population of {temp} rule breaks."

            # Placeholder for recommendation rationale
            line_eight = ""
            if data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] != 0:
                line_eight = "###INSERT TUNING DECISION###"

            # Tuning recommendation
            temp = (
                data_Rule_Pop_Param_Type["Current Threshold"].iloc[0],
                data_Rule_Pop_Param_Type["Recommended Threshold"].iloc[0],
            )

            if param_type_value in currencyF:
                temp_formatted = (f"${temp[0]:,.2f}", f"${temp[1]:,.2f}")
            elif param_type_value in numberF:
                if temp[0] < 10 and temp[1] < 10:
                    temp_formatted = (
                        numbers_df[numbers_df["numbers"] == str(int(temp[0]))][
                            "alpha_numbers"
                        ].iloc[0],
                        numbers_df[numbers_df["numbers"] == str(int(temp[1]))][
                            "alpha_numbers"
                        ].iloc[0],
                    )
                elif temp[0] < 10 and temp[1] >= 10:
                    temp_formatted = (
                        numbers_df[numbers_df["numbers"] == str(int(temp[0]))][
                            "alpha_numbers"
                        ].iloc[0],
                        f"{temp[1]:,}",
                    )
                else:
                    temp_formatted = (f"{temp[0]:,}", f"{temp[1]:,}")
            elif param_type_value in percentF:
                temp_formatted = (f"{temp[0]:.2f}%", f"{temp[1]:.7f}%")
            elif param_type_value in decimalF:
                temp_formatted = (f"{temp[0]:.2f}", f"{temp[1]:.7f}")
            else:
                temp_formatted = (str(temp[0]), str(temp[1]))

            line_nine = ""
            if data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] == 0:
                pass
            elif (
                    data_Rule_Pop_Param_Type["Current Threshold"].iloc[0]
                    == data_Rule_Pop_Param_Type["Recommended Threshold"].iloc[0]
            ):
                line_nine = f"Therefore, it is recommended to maintain the {parameter} threshold at {temp_formatted[0]}."
            elif (
                    data_Rule_Pop_Param_Type["Current Threshold"].iloc[0]
                    != data_Rule_Pop_Param_Type["Recommended Threshold"].iloc[0]
            ):
                line_nine = f"Therefore, it is recommended to lower the {parameter} threshold from {temp_formatted[0]} to {temp_formatted[1]}."

            # Change in effectiveness
            line_ten = ""
            if (
                    data_Rule_Pop_Param_Type["Num Alerts Extracted"].iloc[0] == 0
                    or data_Rule_Pop_Param_Type["Current Threshold"].iloc[0]
                    == data_Rule_Pop_Param_Type["Recommended Threshold"].iloc[0]
            ):
                pass
            elif (
                    data_Rule_Pop_Param_Type["Current Threshold"].iloc[0]
                    != data_Rule_Pop_Param_Type["Recommended Threshold"].iloc[0]
            ):
                line_ten = f"At the recommended threshold, the effectiveness of the delta population is {data_Rule_Pop_Param_Type['Prop Effectiveness'].iloc[0]:.2f}%."

            # Conclusion Paragraph --------------------------------------------------------------
            thresholds_changed = []
            thresholds_changed_values = []
            thresholds_kept = []
            thresholds_kept_values = []

            # Reset index of data_Rule_Pop
            data_Rule_Pop = data_Rule_Pop.reset_index(drop=True)

            # Selects parameter and recommended threshold for each row
            for row in range(len(data_Rule_Pop)):
                temp_param = data_Rule_Pop.iloc[row]["Parameter"]
                temp_val = data_Rule_Pop.iloc[row]["Recommended Threshold"]

                # Formats values as needed
                if temp_param in currencyF:
                    temp_val_formatted = f"${temp_val:,.2f}"
                elif temp_param in numberF:
                    if temp_val < 10:
                        temp_val_formatted = numbers_df[
                            numbers_df["numbers"] == str(int(temp_val))
                            ]["alpha_numbers"].iloc[0]
                    else:
                        temp_val_formatted = f"{temp_val:,}"
                elif temp_param in percentF:
                    temp_val_formatted = f"{temp_val:.7f}%"
                elif temp_param in decimalF:
                    temp_val_formatted = f"{temp_val:.7f}"
                else:
                    temp_val_formatted = str(temp_val)

                # Determines whether the parameter was changed or maintained
                if data_Rule_Pop.iloc[row]["Current Threshold"] != temp_val:
                    thresholds_changed.append(temp_param)
                    thresholds_changed_values.append(temp_val_formatted)
                else:
                    thresholds_kept.append(temp_param)
                    thresholds_kept_values.append(temp_val_formatted)

            # Generate a list of parameters changed with proper formatting
            if len(thresholds_changed) == 1:
                thresholds_changed_formatted = f"{thresholds_changed[0]} parameter"
            elif len(thresholds_changed) > 1:
                thresholds_changed_formatted = f"{', '.join(thresholds_changed[:-1])}, and {thresholds_changed[-1]} parameters"
            else:
                thresholds_changed_formatted = ""

            # Generates a list of thresholds changed with proper formatting
            if len(thresholds_changed_values) == 1:
                thresholds_changed_values_formatted = thresholds_changed_values[0]
            elif len(thresholds_changed_values) > 1:
                thresholds_changed_values_formatted = f"{', '.join(thresholds_changed_values[:-1])}, and {thresholds_changed_values[-1]}"
            else:
                thresholds_changed_values_formatted = ""

            # Generate a list of parameters kept with proper formatting
            if len(thresholds_kept) == 1:
                thresholds_kept_formatted = f"{thresholds_kept[0]} parameter"
            elif len(thresholds_kept) > 1:
                thresholds_kept_formatted = f"{', '.join(thresholds_kept[:-1])}, and {thresholds_kept[-1]} parameters"
            else:
                thresholds_kept_formatted = ""

            # Generates a list of thresholds kept with proper formatting
            if len(thresholds_kept_values) == 1:
                thresholds_kept_values_formatted = thresholds_kept_values[0]
            elif len(thresholds_kept_values) > 1:
                thresholds_kept_values_formatted = f"{', '.join(thresholds_kept_values[:-1])}, and {thresholds_kept_values[-1]}"
            else:
                thresholds_kept_values_formatted = ""

            # Populates the recommendations made
            if len(thresholds_changed) == 0 and len(thresholds_kept) == 0:
                line_twelve = ""
            elif len(thresholds_changed) == 0:
                line_twelve = f"Maintaining the {thresholds_kept_formatted} at {thresholds_kept_values_formatted}."
            elif len(thresholds_kept) == 0:
                line_twelve = f"Adjusting the {thresholds_changed_formatted} to {thresholds_changed_values_formatted}."
            else:
                line_twelve = f"Adjusting the {thresholds_changed_formatted} to {thresholds_changed_values_formatted} while maintaining the {thresholds_kept_formatted} at {thresholds_kept_values_formatted}."

            # Impact on Effectiveness
            temp_eff = data_Rule_Pop_Param_Type["Effectiveness"].iloc[0]
            temp_eff_formatted = f"{temp_eff:.2f}%"

            temp_prop_eff = data_Rule_Pop_Param_Type["Prop Effectiveness"].iloc[0]
            temp_prop_eff_formatted = f"{temp_prop_eff:.2f}%"

            if temp_prop_eff == 0:
                line_thirteen = " respectively will result in no impact to the current alert output in the system."
            else:
                line_thirteen = f"is expected to result in a delta population effectiveness of {temp_prop_eff_formatted}."

            # Expected rule breaks per month
            if temp_prop_eff == 0:
                line_fourteen = ""
            else:
                line_fourteen = "The segment is expected to generate approximately ## additional rule breaks per month."

            # Consolidate Narratives -----------------------------------------------------
            # Populates narratives table made earlier for export
            narratives = pd.concat(
                [
                    narratives,
                    pd.DataFrame(
                        [
                            [
                                data_Rule_Pop_Param_Type["Rule ID"].iloc[0],  # Rule ID
                                data_Rule_Pop_Param_Type["Rule Name"].iloc[
                                    0
                                ],  # Rule Name
                                data_Rule_Pop_Param_Type["Population Group"].iloc[
                                    0
                                ],  # Population group
                                parameter,  # Threshold tuned
                                line_one
                                + " "
                                + line_two
                                + " "
                                + line_three
                                + " "
                                + line_four,  # Tuning summary
                                line_five
                                + " "
                                + line_six
                                + " "
                                + line_seven
                                + " "
                                + line_eight
                                + " "
                                + line_nine
                                + " "
                                + line_ten,  # Analysis
                                line_twelve
                                + " "
                                + line_thirteen
                                + " "
                                + line_fourteen,  # Conclusion
                            ]
                        ],
                        columns=[
                            "Rule ID",
                            "Rule Name",
                            "Population Group",
                            "Parameter",
                            "Summary",
                            "Analysis",
                            "Conclusion",
                        ],
                    ),
                ],
                ignore_index=True,
            )


# Export narratives to CSV
# narratives.to_csv(
#     os.path.join(fp, "Actimize Pre-Production Report Narratives.csv"), index=False
# )

# Export Narratives to Word ----------------------------------------

# Initialize output file
output = os.path.join(
    r"C:/Users/KadamatiV/OneDrive - Crowe LLP/Documents/PROJECTHUB/TM TUNING/TM_Narrative_Codebase_Consolidation/output/",
    "CNB Report Narratives - BTL latest_512024_345pm.docx",
)

doc = docx.Document()

# Disclaimer for edit
# Modify the built-in 'Heading 1' style
heading_style = doc.styles["Heading 1"]
heading_style.font.name = "Times New Roman"
heading_style.font.size = Pt(10)
heading_style.font.bold = True
heading_style.font.color.rgb = RGBColor(255, 0, 0)  # Set font color to red

# Use the modified 'Heading 1' style for your heading
heading = doc.add_heading("", level=1)
run = heading.add_run(
    'Manual edit is required for the following instances of "###INSERT TUNING DECISION###"'
)
run.font.name = "Times New Roman"  # Set the font of the run to Times New Roman
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)

# Iterate through each paragraph in the document
for paragraph in doc.paragraphs:
    # Iterate through each run in the paragraph
    for run in paragraph.runs:
        # Check if the specific text is in the run
        if "###INSERT TUNING DECISION###" in run.text:
            # Change the font color of the run to red
            run.font.color.rgb = RGBColor(255, 0, 0)

# ------ end of disclaimer + search utility for specific string and color aspect


# ? Formatting Configurations
# Create a new style based on 'Heading 2'
new_heading_style = doc.styles.add_style("NewHeading2", WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = doc.styles["Heading 2"]
new_heading_style.font.color.rgb = RGBColor(0, 0, 0)
new_heading_style.font.name = "Times New Roman"
new_heading_style.font.size = Pt(10)
new_heading_style.font.bold = True

# pop formatting
# Modify the built-in 'Heading 1' style
heading_style_3 = doc.styles["Heading 1"]
heading_style_3.font.name = "Times New Roman"
heading_style_3.font.size = Pt(10)
heading_style_3.font.bold = True

# Use the modified 'Heading 1' style for your heading
heading = doc.add_heading("", level=1)  # Add an empty heading
run = heading.add_run("Below-the-Line Analysis")  # Add text to the heading with a run
run.font.name = "Times New Roman"  # Set the font of the run to Times New Roman
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)

# Filter narratives to rule level
for x in natsorted(ruleIDs):
    # for x in natsorted(ruleIDs):
    narratives_Rule = narratives[narratives["Rule ID"] == x]
    # Add Rule ID header to doc
    # heading = doc.add_heading(x, level=2)
    # Add Rule ID header to doc using the new style
    heading = doc.add_heading(x, level=2)
    heading.style = new_heading_style

    # Add threshold decisions table
    # ! change log v1 // formatted the above as a paragraph
    # Check if the style 'SummaryOfThresholdDecisions' exists, and if not, add it
    style_name = "SummaryOfThresholdDecisions"
    if not style_name in doc.styles:
        summary_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        summary_style.font.name = "Times New Roman"
        summary_style.font.size = Pt(10)
        summary_style.font.bold = False
        summary_style.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
    else:
        summary_style = doc.styles[style_name]
        summary_style.font.name = "Times New Roman"
        summary_style.font.size = Pt(10)
        summary_style.font.bold = False
        summary_style.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

    # Add a paragraph with the defined style
    paragraph = doc.add_paragraph("Summary of Threshold Decisions", style=style_name)

    threshold_decisions = data[
        [
            "Rule ID",
            "Population Group",
            "Parameter",
            "Current Threshold",
            "BTL Threshold",
            "Recommended Threshold",
        ]
    ]
    threshold_decisions = threshold_decisions[threshold_decisions["Rule ID"] == x]
    threshold_decisions = threshold_decisions[
        [
            "Population Group",
            "Parameter",
            "Current Threshold",
            "BTL Threshold",
            "Recommended Threshold",
        ]
    ]
    threshold_decisions = threshold_decisions.sort_values(
        by=["Population Group", "Parameter"]
    )

    threshold_decisions = threshold_decisions.rename(
        columns={"Population Group": "Population"}
    )

    # Format thresholds in table to display with proper units
    for row in threshold_decisions.index:
        temp_param = threshold_decisions.at[row, "Parameter"]
        temp_values = (
            threshold_decisions.loc[
                row, ["Current Threshold", "BTL Threshold", "Recommended Threshold"]
            ]
            .astype(str)
            .values
        )

        if temp_param in currencyF:
            temp_values_formatted = [
                f"${float(value.replace(',', '')):,.2f}" for value in temp_values
            ]
        elif temp_param in numberF:
            temp_values_formatted = [f"{float(value):,.0f}" for value in temp_values]
        elif temp_param in percentF:
            temp_values_formatted = [
                f"{float(value.strip('%')) / 100:.7%}" for value in temp_values
            ]
        elif temp_param in decimalF:
            temp_values_formatted = [f"{float(value):.2f}" for value in temp_values]
        else:
            temp_values_formatted = [str(value) for value in temp_values]

        threshold_decisions.at[row, "Current Threshold"] = temp_values_formatted[0]
        threshold_decisions.at[row, "BTL Threshold"] = temp_values_formatted[1]
        threshold_decisions.at[row, "Recommended Threshold"] = temp_values_formatted[2]

    # ! ChangeLog v1: table formatting
    # Create a table
    table = doc.add_table(rows=1, cols=len(threshold_decisions.columns))

    # Define the header cells
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(threshold_decisions.columns):
        hdr_cells[i].text = col_name
        hdr_cells[i].vertical_alignment = (
            WD_ALIGN_VERTICAL.BOTTOM
        )  # Align text to bottom (optional)
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True  # Make header text bold

    # Add rows to the table
    for _, row in threshold_decisions.iterrows():
        row_cells = table.add_row().cells
        for i, col_value in enumerate(row):
            row_cells[i].text = str(col_value)
            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


    # Function to set cell borders
    def set_cell_borders(cell, border_type, sz=6, color="auto"):
        """
        Set cell borders. Specify the border_type as 'top', 'bottom', 'start', or 'end'.
        Set 'sz' to the size of the border and 'color' to the border color.
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tag = "w:{}".format(border_type)
        element = tcPr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcPr.append(element)
        element.set(qn("w:sz"), str(sz))
        element.set(qn("w:val"), "single")
        if color != "auto":
            element.set(qn("w:color"), color)


    # Apply formatting to the table
    for row in table.rows:
        for cell in row.cells:
            # Clear all cell borders
            set_cell_borders(cell, "top", sz=0, color="FFFFFF")
            set_cell_borders(cell, "bottom", sz=0, color="FFFFFF")
            set_cell_borders(cell, "start", sz=0, color="FFFFFF")
            set_cell_borders(cell, "end", sz=0, color="FFFFFF")

    # Set top and bottom borders for header row
    for cell in table.rows[0].cells:
        set_cell_borders(cell, "top", sz=6, color="000000")
        set_cell_borders(cell, "bottom", sz=6, color="000000")

    # Set bottom border for last row
    for cell in table.rows[-1].cells:
        set_cell_borders(cell, "bottom", sz=6, color="000000")

    doc.add_paragraph("")

    # Filter to Pop Group
    popGroups = narratives_Rule["Population Group"].value_counts().index.tolist()

    for pop in popGroups:
        narratives_Rule_Pop = narratives_Rule[
            narratives_Rule["Population Group"] == pop
            ]

        # ! Chnagelog v1: threshold recommendation formatting
        heading_style_4 = doc.styles["Heading 4"]
        heading_style_4.font.name = "Times New Roman"
        heading_style_3.font.size = Pt(10)
        heading_style_3.font.bold = True

        # Use the modified 'Heading 1' style for your heading
        heading = doc.add_heading("", level=4)  # Add an empty heading
        run = heading.add_run(pop)  # Add text to the heading with a run
        run.font.name = "Times New Roman"  # Set the font of the run to Times New Roman
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Add header for population group
        # heading = doc.add_heading(pop, level=3)
        # ! Chnagelog v1: pop group formatting
        # heading_style_3 = doc.styles["Heading 3"]
        heading_style_3.font.name = "Times New Roman"
        heading_style_3.font.size = Pt(10)
        heading_style_3.font.bold = True

        # Use the modified 'Heading 1' style for your heading
        heading = doc.add_heading("", level=3)  # Add an empty heading
        run = heading.add_run(
            "Threshold Recommendation"
        )  # Add text to the heading with a run
        run.font.name = "Times New Roman"  # Set the font of the run to Times New Roman
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Add 'Threshold Recommendation' header to doc
        # heading = doc.add_heading("Threshold Recommendation", level=4)

        # Print summary paragraph
        # paragraph = doc.add_paragraph(narratives_Rule_Pop["Summary"].iloc[0])
        # ! change log v1 // formatted the above as a paragraph
        # Check if the style 'SummaryOfThresholdDecisions' exists, and if not, add it
        style_name_narrative_rule_pop_summary = "narrative_rule_pop_summary"
        if not style_name_narrative_rule_pop_summary in doc.styles:
            summary_style = doc.styles.add_style(
                style_name_narrative_rule_pop_summary, WD_STYLE_TYPE.PARAGRAPH
            )
            summary_style.font.name = "Times New Roman"
            summary_style.font.size = Pt(10)
            summary_style.font.bold = False
            summary_style.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
        else:
            summary_style = doc.styles[style_name_narrative_rule_pop_summary]
            summary_style.font.name = "Times New Roman"
            summary_style.font.size = Pt(10)
            summary_style.font.bold = False
            summary_style.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

        # Add a paragraph with the defined style
        paragraph = doc.add_paragraph(
            narratives_Rule_Pop["Summary"].iloc[0],
            style=style_name_narrative_rule_pop_summary,
        )

        # Skip remaining headers/paragraphs if no alerts generated
        if narratives_Rule_Pop["Summary"].iloc[0].startswith("No alerts"):
            continue

        # Filter to each parameter within pop group
        params = narratives_Rule_Pop["Parameter"].value_counts().index.tolist()

        for param in params:
            narratives_Rule_Pop_Parameter = narratives_Rule_Pop[
                narratives_Rule_Pop["Parameter"] == param
                ]

            # Add parameter header
            # heading = doc.add_heading(param, level=5)
            # ! Chnagelog v1: param formatting
            heading_style_5 = doc.styles["Heading 5"]
            heading_style_5.font.name = "Times New Roman"
            heading_style_5.font.size = Pt(10)
            heading_style_5.font.bold = True

            # Use the modified 'Heading 1' style for your heading
            heading = doc.add_heading("", level=5)  # Add an empty heading

            run = heading.add_run(param)  # Add text to the heading with a run
            run.font.name = (
                "Times New Roman"  # Set the font of the run to Times New Roman
            )
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Add analysis paragraphs for each parameter
            # paragraph = doc.add_paragraph(
            #     narratives_Rule_Pop_Parameter["Analysis"].iloc[0]
            # )
            # ! change log v1 // formatted the above as a paragraph
            # Check if the style 'SummaryOfThresholdDecisions' exists, and if not, add it
            style_name_narrative_rule_pop_parameter = "narrative_rule_pop_parameter"
            if not style_name_narrative_rule_pop_parameter in doc.styles:
                summary_style = doc.styles.add_style(
                    style_name_narrative_rule_pop_parameter, WD_STYLE_TYPE.PARAGRAPH
                )
                summary_style.font.name = "Times New Roman"
                summary_style.font.size = Pt(10)
                summary_style.font.bold = False
                summary_style.font.color.rgb = RGBColor(
                    0, 0, 0
                )  # Set font color to black
            else:
                summary_style = doc.styles[style_name_narrative_rule_pop_parameter]
                summary_style.font.name = "Times New Roman"
                summary_style.font.size = Pt(10)
                summary_style.font.bold = False
                summary_style.font.color.rgb = RGBColor(
                    0, 0, 0
                )  # Set font color to black

            # Add a paragraph with the defined style
            paragraph = doc.add_paragraph(
                narratives_Rule_Pop_Parameter["Analysis"].iloc[0],
                style=style_name_narrative_rule_pop_parameter,
            )

            # Get the original and lowered threshold values for the current parameter
            temp_original = data_Rule_Pop_Parameter["Current Threshold"].iloc[0]
            temp_lowered = data_Rule_Pop_Parameter["Recommended Threshold"].iloc[0]

            if param in currencyF:
                temp_original_formatted = f"${temp_original:,.2f}"
                temp_lowered_formatted = f"${temp_lowered:,.2f}"
            elif param in numberF:
                if temp_original < 10:
                    temp_original_formatted = numbers_df[
                        numbers_df["numbers"] == str(int(temp_original))
                        ]["alpha_numbers"].iloc[0]
                else:
                    temp_original_formatted = f"{temp_original:,}"

                if temp_lowered < 10:
                    temp_lowered_formatted = numbers_df[
                        numbers_df["numbers"] == str(int(temp_lowered))
                        ]["alpha_numbers"].iloc[0]
                else:
                    temp_lowered_formatted = f"{temp_lowered:,}"
            elif param in percentF:
                temp_original_formatted = f"{temp_original:.7f}%"
                temp_lowered_formatted = f"{temp_lowered:.7f}%"
            elif param in decimalF:
                temp_original_formatted = f"{temp_original:.2f}"
                temp_lowered_formatted = f"{temp_lowered:.2f}"
            else:
                temp_original_formatted = str(temp_original)
                temp_lowered_formatted = str(temp_lowered)

            # Add conclusion paragraph for each parameter
            conclusion = f"Based on the analysis, it is recommended to lower the {param} threshold from {temp_original_formatted} to {temp_lowered_formatted}."
            # paragraph = doc.add_paragraph(conclusion)

        # Add header for conclusion
        # heading = doc.add_heading("Conclusion", level=4)
        # ! Changelog v1: conclusion heading format
        heading_style_4 = doc.styles["Heading 4"]
        heading_style_4.font.name = "Times New Roman"
        heading_style_4.font.size = Pt(10)
        heading_style_4.font.bold = True
        heading_style_4.font.italic = False
        # Use the modified 'Heading 1' style for your heading
        heading = doc.add_heading("", level=4)  # Add an empty heading
        run = heading.add_run("Conclusion")  # Add text to the heading with a run
        run.font.name = "Times New Roman"  # Set the font of the run to Times New Roman
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = False
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Print conclusion paragraph
        conclusion = narratives_Rule_Pop["Conclusion"].iloc[0]
        # paragraph = doc.add_paragraph(conclusion)

        # ! change log v1 // formatted the above as a paragraph : conclusions
        # Check if the style 'SummaryOfThresholdDecisions' exists, and if not, add it
        style_name_conclusion = "style_name_conclusions"
        if not style_name_conclusion in doc.styles:
            summary_style = doc.styles.add_style(
                style_name_conclusion, WD_STYLE_TYPE.PARAGRAPH
            )
            summary_style.font.name = "Times New Roman"
            summary_style.font.size = Pt(10)
            summary_style.font.bold = False
            summary_style.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
        else:
            summary_style = doc.styles[style_name_conclusion]
            summary_style.font.name = "Times New Roman"
            summary_style.font.size = Pt(10)
            summary_style.font.bold = False
            summary_style.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

        # Add a paragraph with the defined style
        paragraph = doc.add_paragraph(
            narratives_Rule_Pop["Conclusion"].iloc[0],
            style=style_name_conclusion,
        )


def color_text_red(doc):
    target_text = "###INSERT TUNING DECISION###"
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if target_text in run.text:
                # Split text to isolate the target text for coloring
                texts = run.text.split(target_text)
                run.text = texts[0]  # Text before the target text
                for text in texts[1:]:
                    # Add the target text and set its color to red
                    colored_run = paragraph.add_run(target_text)
                    colored_run.font.color.rgb = RGBColor(255, 0, 0)
                    # Add the rest of the text after the target text
                    normal_run = paragraph.add_run(text)
                    normal_run.font.color.rgb = RGBColor(
                        0, 0, 0
                    )  # Assuming default color is black


# Apply the function to color text
color_text_red(doc)
doc.save(output)
