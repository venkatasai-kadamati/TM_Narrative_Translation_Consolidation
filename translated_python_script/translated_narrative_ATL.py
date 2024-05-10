import pandas as pd
import numpy as np
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Specify location of the tuning tracker
fp = "C:/Users/KadamatiV/OneDrive - Crowe LLP/Documents/PROJECTHUB/TM TUNING/TM_Narrative_Codebase_Consolidation/required_processing_data/"
file_name = "Tuning Tracker - ATL Calculationsv11.xlsx"

report_name = "output/CNB - Production Report Narratives ATL new.docx"

# Specify how to format the different parameters seen (change threshold names as needed)
currencyF = [
    "Minimal Sum",
    "Minimum Value",
    "Minimal Transaction Amount",
    "Sum Lower Bound",
    "Min Value",
    "Minimal Current Month Sum",
    "Minimal Transaction Value",
    "Transaction Amount Lower Bound",
]
numberF = ["No. of Occurrences", "Minimum Volume", "Min Value"]
percentF = ["Ratio Lower Bound", "Ratio Upper Bound"]
decimalF = [
    "STDEV exceeds Historical Average Sum",
    "STDEV exceeds Historical Average Count",
]


def format_threshold(value, threshold_type):
    if threshold_type in currencyF:
        return f"${value:,.2f}"
    elif threshold_type in numberF:
        return f"{value:,}"
    elif threshold_type in percentF:
        return f"{value * 100:.2f}%"
    elif threshold_type in decimalF:
        return f"{value:,.2f}"
    else:
        return str(value)


# Data is loaded
data = pd.read_excel(fp + file_name, sheet_name=0)

# Enter the column name of the Rule IDs, if different
ruleIDs = data["Rule ID"].value_counts().index

# Create lookup for values below 10 that need to be written out (no change)
numbers = {
    "0": "zero (0)",
    "1": "one (1)",
    "2": "two (2)",
    "3": "three (3)",
    "4": "four (4)",
    "5": "five (5)",
    "6": "six (6)",
    "7": "seven (7)",
    "8": "eight (8)",
    "9": "nine (9)",
}
numbers_cap = {k: v.capitalize() for k, v in numbers.items()}

# Create an empty data frame to hold narratives
narratives = pd.DataFrame(
    columns=[
        "Rule ID",
        "Rule Name",
        "Population Group",
        "Parameter",
        "Summary",
        "Analysis",
        "Conclusion",
    ]
)

# Begin Iterations --------------------------------------------------------

for x in ruleIDs:
    data_Rule = data[data["Rule ID"] == x]
    popGroups = data_Rule["Population Group"].value_counts().index

    for pop in popGroups:
        data_Rule_Pop = data_Rule[data_Rule["Population Group"] == pop]
        params = data_Rule_Pop["Parameter"].value_counts().index

        for threshold in params:
            data_Rule_Pop_Parameter = data_Rule_Pop[
                data_Rule_Pop["Parameter"] == threshold
            ]

            # Summary Paragraph -------------------------------------------------------
            temp = data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0]
            date_range = data_Rule_Pop_Parameter["Date Range"].values[0].split("-")
            params_used = ", ".join(data_Rule_Pop_Parameter["Parameter"].unique())

            line_one = (
                "No alerts generated"
                if temp == 0
                else f"Production tuning analysis was performed on the {params_used} for the {data_Rule_Pop['Population Group'].values[0]} population group."
            )

            line_two = (
                ""
                if temp == 0
                else (
                    f"{numbers[str(temp)]} rule break was generated"
                    if temp == 1
                    else f"{temp} rule breaks were generated"
                )
                + f" between {date_range[0]} and {date_range[1]}."
            )

            line_three = (
                ""
                if temp == 0
                else (
                    f"The Bank identified {numbers[str(temp)]} Data Quality rule break"
                    if temp == 1
                    else f"The Bank identified {temp} Data Quality rule breaks"
                )
                + " that generated on duplicated or incorrectly mapped transactional activity. As a result, these rule breaks were marked as Data Quality and excluded from analysis."
            )

            temp_total = temp - data_Rule_Pop_Parameter["Data Quality Alerts"].values[0]
            temp_sar = data_Rule_Pop_Parameter["SARs Filed"].values[0]
            temp_int = (
                data_Rule_Pop_Parameter["Interesting Alerts"].values[0] + temp_sar
            )
            temp_eff = round(data_Rule_Pop_Parameter["Effectiveness"].values[0], 2)
            temp_sar_yield = round(data_Rule_Pop_Parameter["SAR Yield"].values[0], 2)

            line_four = (
                ""
                if temp_total == 0
                else (
                    f"The one (1) reviewed rule break was determined to be Interesting and led to a SAR filing, resulting in a production effectiveness and SAR yield of 100.00%."
                    if temp_total == 1 and temp_int == 1 and temp_sar == 1
                    else (
                        f"The one (1) reviewed rule break was determined to be Interesting but did not end in a SAR filing, resulting in a production effectiveness of 100.00% and a SAR yield of 0.00%."
                        if temp_total == 1 and temp_int == 1 and temp_sar == 0
                        else (
                            f"The one (1) reviewed rule break was not determined to be Interesting, resulting in a production effectiveness of 0.00%."
                            if temp_total == 1 and temp_int != 1
                            else (
                                f"Of the total {numbers[str(temp_total)] if temp_total < 10 else temp_total} reviewed rule breaks, {numbers[str(temp_int)] if temp_int < 10 else temp_int} rule break was determined to be Interesting, resulting in a production effectiveness of {temp_eff}%."
                                if temp_total < 10 and temp_int == 1
                                else f"Of the total {numbers[str(temp_total)] if temp_total < 10 else temp_total} reviewed rule breaks, {numbers[str(temp_int)] if temp_int < 10 else temp_int} rule breaks were determined to be Interesting, resulting in a production effectiveness of {temp_eff}%."
                            )
                        )
                    )
                )
            )

            line_five = (
                ""
                if temp_total == 0 or temp_int == 0
                else (
                    f"{numbers[str(temp_sar)]} of the {numbers[str(temp_int)] if temp_int < 10 else temp_int} Interesting rule breaks led to a SAR filing, resulting in a SAR yield of {temp_sar_yield}%. Additional detail on the analysis results per population group is provided below."
                    if temp_sar < 10 and temp_int < 10
                    else (
                        f"{numbers[str(temp_sar)]} of the {temp_int} Interesting rule breaks led to a SAR filing, resulting in a SAR yield of {temp_sar_yield}%. Additional detail on the analysis results per population group is provided below."
                        if temp_sar < 10 and temp_int >= 10
                        else f"{temp_sar} of the {temp_int} Interesting rule breaks led to a SAR filing, resulting in a SAR yield of {temp_sar_yield}%. Additional detail on the analysis results per population group is provided below."
                    )
                )
            )

            # Analysis Paragraph ------------------------------------------------------
            temp = data_Rule_Pop_Parameter["Current Threshold"].values[0]
            if threshold in currencyF:
                temp_formatted = f"${temp:,.2f}"
            elif threshold in numberF:
                temp_formatted = (
                    numbers[str(int(temp))]
                    if temp < 10 and temp % 1 == 0
                    else f"{temp:,}"
                )

            elif threshold in percentF:
                temp_formatted = f"{temp:.0%}"
            elif threshold in decimalF:
                temp_formatted = f"{temp:,.2f}"
            else:
                temp_formatted = temp

            line_six = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                else f"Above-the-line tuning was conducted on the {data_Rule_Pop_Parameter['Parameter'].values[0]} threshold, which was set at a value of {temp_formatted}."
            )

            temp = (
                data_Rule_Pop_Parameter["Max Val"].values[0]
                - data_Rule_Pop_Parameter["Min Val"].values[0]
            )
            val_formatted = []
            if threshold in currencyF:
                val_formatted = [
                    f"${val:,.2f}"
                    for val in [
                        data_Rule_Pop_Parameter["Min Val"].values[0],
                        data_Rule_Pop_Parameter["Max Val"].values[0],
                    ]
                ]
            elif threshold in numberF:
                val_formatted = [
                    numbers[str(int(val))] if val < 10 and val % 1 == 0 else f"{val:,}"
                    for val in [
                        data_Rule_Pop_Parameter["Min Val"].values[0],
                        data_Rule_Pop_Parameter["Max Val"].values[0],
                    ]
                ]

            elif threshold in percentF:
                val_formatted = [
                    f"{val:.0%}"
                    for val in [
                        data_Rule_Pop_Parameter["Min Val"].values[0],
                        data_Rule_Pop_Parameter["Max Val"].values[0],
                    ]
                ]
            elif threshold in decimalF:
                val_formatted = [
                    f"{val:,.2f}"
                    for val in [
                        data_Rule_Pop_Parameter["Min Val"].values[0],
                        data_Rule_Pop_Parameter["Max Val"].values[0],
                    ]
                ]
            else:
                val_formatted = [
                    data_Rule_Pop_Parameter["Min Val"].values[0],
                    data_Rule_Pop_Parameter["Max Val"].values[0],
                ]

            line_seven = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                else (
                    f"Rule breaks were generated solely at a value of {val_formatted[1]} within the {data_Rule_Pop['Population Group'].values[0]} population segment."
                    if temp == 0
                    else f"Rule breaks were generated for values ranging between {val_formatted[0]} and {val_formatted[1]} within the {data_Rule_Pop['Population Group'].values[0]} population segment."
                )
            )

            line_eight = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                else (
                    f"One (1) Interesting rule break was noted in the production population which also resulted in a SAR filing."
                    if temp_int == 1 and temp_sar == 1
                    else (
                        f"One (1) Interesting rule break was noted in the production population which did not result in a SAR filing."
                        if temp_int == 1 and temp_sar != 1
                        else (
                            f"{numbers[str(temp_int)]} Interesting rule breaks were noted in the production population, of which one (1) rule break resulted in a SAR filing."
                            if temp_int < 10 and temp_sar == 1
                            else (
                                f"{numbers[str(temp_int)]} Interesting rule breaks were noted in the production population, of which {numbers[str(temp_sar)]} rule breaks resulted in SAR filings."
                                if temp_int < 10
                                else (
                                    f"{temp_int} Interesting rule breaks were noted in the production population, of which one (1) rule break resulted in a SAR filing."
                                    if temp_int >= 10 and temp_sar == 1
                                    else (
                                        f"{temp_int} Interesting rule breaks were noted in the production population, of which {numbers[str(temp_sar)]} rule breaks resulted in SAR filings."
                                        if temp_int >= 10 and temp_sar < 10
                                        else f"{temp_int} Interesting rule breaks were noted in the production population, of which {temp_sar} rule breaks resulted in SAR filings."
                                    )
                                )
                            )
                        )
                    )
                )
            )

            line_nine = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                else "###INSERT TUNING DECISION###"
            )

            temp = [
                data_Rule_Pop_Parameter["Current Threshold"].values[0],
                data_Rule_Pop_Parameter["Recommended Threshold"].values[0],
            ]
            if threshold in currencyF:
                temp_formatted = [f"${val:,.2f}" for val in temp]
            elif threshold in numberF:
                temp_formatted = [
                    numbers[str(int(val))] if val < 10 and val % 1 == 0 else f"{val:,}"
                    for val in temp
                ]

            elif threshold in percentF:
                temp_formatted = [f"{val:.0%}" for val in temp]
            elif threshold in decimalF:
                temp_formatted = [f"{val:,.2f}" for val in temp]
            else:
                temp_formatted = temp

            line_ten = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                else (
                    f"Therefore, it is recommended to maintain the {data_Rule_Pop_Parameter['Parameter'].values[0]} threshold at {temp_formatted[0]}."
                    if data_Rule_Pop_Parameter["Current Threshold"].values[0]
                    == data_Rule_Pop_Parameter["Recommended Threshold"].values[0]
                    else f"Therefore, it is recommended to increase the {data_Rule_Pop_Parameter['Parameter'].values[0]} threshold from {temp_formatted[0]} to {temp_formatted[1]}."
                )
            )

            line_eleven = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                or data_Rule_Pop_Parameter["Effectiveness"].values[0]
                > data_Rule_Pop_Parameter["Prop Effectiveness"].values[0]
                else (
                    f"At the recommended threshold, the overall effectiveness will remain at {data_Rule_Pop_Parameter['Effectiveness'].values[0]:.2f}%."
                    if data_Rule_Pop_Parameter["SAR Yield"].values[0]
                    > data_Rule_Pop_Parameter["Prop SAR Yield"].values[0]
                    and (
                        data_Rule_Pop_Parameter["Current Threshold"].values[0]
                        == data_Rule_Pop_Parameter["Recommended Threshold"].values[0]
                        or data_Rule_Pop_Parameter["Effectiveness"].values[0]
                        == data_Rule_Pop_Parameter["Prop Effectiveness"].values[0]
                    )
                    else (
                        f"At the recommended threshold, the overall effectiveness will remain at {data_Rule_Pop_Parameter['Effectiveness'].values[0]:.2f}%"
                        if data_Rule_Pop_Parameter["Current Threshold"].values[0]
                        == data_Rule_Pop_Parameter["Recommended Threshold"].values[0]
                        or data_Rule_Pop_Parameter["Effectiveness"].values[0]
                        == data_Rule_Pop_Parameter["Prop Effectiveness"].values[0]
                        else (
                            f"At the recommended threshold, the overall effectiveness will increase from {data_Rule_Pop_Parameter['Effectiveness'].values[0]:.2f}% to {data_Rule_Pop_Parameter['Prop Effectiveness'].values[0]:.2f}%."
                            if data_Rule_Pop_Parameter["Effectiveness"].values[0]
                            < data_Rule_Pop_Parameter["Prop Effectiveness"].values[0]
                            and data_Rule_Pop_Parameter["SAR Yield"].values[0]
                            > data_Rule_Pop_Parameter["Prop SAR Yield"].values[0]
                            else f"At the recommended threshold, the overall effectiveness will increase from {data_Rule_Pop_Parameter['Effectiveness'].values[0]:.2f}% to {data_Rule_Pop_Parameter['Prop Effectiveness'].values[0]:.2f}%"
                        )
                    )
                )
            )

            line_twelve = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                or data_Rule_Pop_Parameter["SAR Yield"].values[0]
                > data_Rule_Pop_Parameter["Prop SAR Yield"].values[0]
                else (
                    f"At the recommended threshold, the overall SAR yield will remain at {data_Rule_Pop_Parameter['SAR Yield'].values[0]:.2f}%."
                    if line_eleven == ""
                    and (
                        data_Rule_Pop_Parameter["Current Threshold"].values[0]
                        == data_Rule_Pop_Parameter["Recommended Threshold"].values[0]
                        or data_Rule_Pop_Parameter["SAR Yield"].values[0]
                        == data_Rule_Pop_Parameter["Prop SAR Yield"].values[0]
                    )
                    else (
                        "and the overall SAR yield will remain at {data_Rule_Pop_Parameter['SAR Yield'].values[0]:.2f}%."
                        if data_Rule_Pop_Parameter["Current Threshold"].values[0]
                        == data_Rule_Pop_Parameter["Recommended Threshold"].values[0]
                        or data_Rule_Pop_Parameter["SAR Yield"].values[0]
                        == data_Rule_Pop_Parameter["Prop SAR Yield"].values[0]
                        else (
                            f"At the recommended threshold, the overall SAR yield will increase from {data_Rule_Pop_Parameter['SAR Yield'].values[0]:.2f}% to {data_Rule_Pop_Parameter['Prop SAR Yield'].values[0]:.2f}%."
                            if line_eleven == ""
                            and data_Rule_Pop_Parameter["SAR Yield"].values[0]
                            < data_Rule_Pop_Parameter["Prop SAR Yield"].values[0]
                            else f"and the overall SAR yield will increase from {data_Rule_Pop_Parameter['SAR Yield'].values[0]:.2f}% to {data_Rule_Pop_Parameter['Prop SAR Yield'].values[0]:.2f}%."
                        )
                    )
                )
            )

            line_thirteen = (
                ""
                if data_Rule_Pop_Parameter["Num Alerts Extracted"].values[0] == 0
                or data_Rule_Pop_Parameter["Not Interesting Alert Reduction"].values[0]
                == 0
                else (
                    f"The recommended threshold will reduce the number of not interesting rule breaks by approximately {data_Rule_Pop_Parameter['Not Interesting Alert Reduction'].values[0]:.2f}%."
                    if data_Rule_Pop_Parameter["Effectiveness"].values[0]
                    > data_Rule_Pop_Parameter["Prop Effectiveness"].values[0]
                    else "Additionally, the recommended threshold will reduce the number of not interesting rule breaks by approximately {data_Rule_Pop_Parameter['Not Interesting Alert Reduction'].values[0]:.2f}%."
                )
            )

            # Consolidate Narratives --------------------------------------------------
            new_data = pd.DataFrame(
                {
                    "Rule ID": data_Rule_Pop_Parameter["Rule ID"].values[0],
                    "Rule Name": data_Rule_Pop_Parameter["Rule Name"].values[0],
                    "Population Group": data_Rule_Pop_Parameter[
                        "Population Group"
                    ].values[0],
                    "Parameter": data_Rule_Pop_Parameter["Parameter"].values[0],
                    "Summary": " ".join(
                        [line_one, line_two, line_three, line_four, line_five]
                    ).strip(),
                    "Analysis": " ".join(
                        [
                            line_six,
                            line_seven,
                            line_eight,
                            line_nine,
                            line_ten,
                            line_eleven,
                            line_twelve,
                            line_thirteen,
                        ]
                    ).strip(),
                    "Conclusion": "",
                },
                index=[0],
            )
            narratives = pd.concat([narratives, new_data], ignore_index=True)

# Export Narratives to Word -----------------------------------------------

document = Document()

document.add_heading("Analysis", level=1)

for x in ruleIDs:
    narratives_Rule = narratives[narratives["Rule ID"] == x]

    rule_name = narratives_Rule["Rule Name"].values[0]
    rule_id = narratives_Rule["Rule ID"].values[0]

    document.add_heading(f"{rule_name} | {rule_id}", level=2)

    threshold_decisions = data[
        ["Population Group", "Parameter", "Current Threshold", "Recommended Threshold"]
    ][data["Rule ID"] == x].sort_values(by=["Population Group", "Parameter"])
    threshold_decisions.columns = [
        "Population",
        "Parameter",
        "Current Threshold",
        "Recommended Threshold",
    ]

    document.add_paragraph("Summary of Threshold Decisions")

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Population"
    header_cells[1].text = "Parameter"
    header_cells[2].text = "Current Threshold"
    header_cells[3].text = "Recommended Threshold"

    for _, row in threshold_decisions.iterrows():
        cells = table.add_row().cells
        cells[0].text = row["Population"]
        cells[1].text = row["Parameter"]
        cells[2].text = format_threshold(row["Current Threshold"], row["Parameter"])
        cells[3].text = format_threshold(row["Recommended Threshold"], row["Parameter"])

    document.add_paragraph("")

    popGroups = narratives_Rule["Population Group"].unique()

    for pop in popGroups:
        narratives_Rule_Pop = narratives_Rule[
            narratives_Rule["Population Group"] == pop
        ]

        document.add_heading(pop, level=3)

        document.add_heading("Threshold Recommendation", level=4)

        document.add_paragraph(narratives_Rule_Pop["Summary"].values[0])

        if "No alerts" not in narratives_Rule_Pop["Summary"].values[0]:
            params = narratives_Rule_Pop["Parameter"].unique()

            for param in params:
                narratives_Rule_Pop_Parameter = narratives_Rule_Pop[
                    narratives_Rule_Pop["Parameter"] == param
                ]

                document.add_heading(param, level=5)

                document.add_paragraph(
                    narratives_Rule_Pop_Parameter["Analysis"].values[0]
                )

            document.add_heading("Conclusion", level=4)

            document.add_paragraph(narratives_Rule_Pop["Conclusion"].values[0])

    document.add_heading("Scoring Recommendation", level=3)

document.save(fp + report_name)
